import logging
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
import chardet
import csv
import fitz  # PyMuPDF
import io
from docx import Document as DocxDocument
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone
from dotenv import load_dotenv
import re
import time
import boto3
from langchain.docstore.document import Document as LangDocument
from langchain.text_splitter import Language, RecursiveCharacterTextSplitter
from embeddings import get_embeddings
from constants import DOCUMENT_MAP, INGEST_THREADS, SOURCE_DIRECTORY

load_dotenv()
pinecone_api_key = os.environ.get('PINECONE_API_KEY')
pc = Pinecone(api_key=pinecone_api_key)

def file_log(logentry):
    print(logentry + "\n")
    return

def detect_encoding(file_content: bytes) -> str:
    encoding_result = chardet.detect(file_content)
    return encoding_result['encoding']

def load_single_document(file_name: str, file_content: bytes) -> LangDocument:
    try:
        file_log(f"Processing file: {file_name}")
        file_extension = os.path.splitext(file_name)[1]
        if file_extension == '.pdf':
            content = read_pdf(file_content)
        elif file_extension == '.docx':
            content = read_docx(file_content)
        elif file_extension == '.csv':
            encoding = detect_encoding(file_content)
            content = read_csv(file_content, encoding)
        else:
            file_log(f"{file_name} document type is undefined.")
            raise ValueError("Document type is undefined")
        return LangDocument(content, metadata={"source": file_name})
    except Exception as ex:
        file_log(f"{file_name} loading error: \n{ex}")
        return None

def read_pdf(file_content: bytes) -> str:
    doc = fitz.open(stream=file_content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def read_docx(file_content: bytes) -> str:
    file_stream = io.BytesIO(file_content)
    doc = DocxDocument(file_stream)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_csv(file_content: bytes, encoding: str) -> str:
    content = file_content.decode(encoding)
    sample = content[:10000]
    dialect = csv.Sniffer().sniff(sample)
    delimiter = dialect.delimiter
    reader = csv.reader(content.splitlines(), delimiter=delimiter)
    fieldnames = next(reader)
    rows = [row for row in reader]
    return "\n".join([delimiter.join(fieldnames)] + [delimiter.join(row) for row in rows])

def load_documents_from_s3(bucket_name: str, folder: str) -> list[LangDocument]:
    s3 = boto3.client('s3')
    paginator = s3.get_paginator('list_objects_v2')
    pages = paginator.paginate(Bucket=bucket_name, Prefix=folder)

    files = []
    for page in pages:
        for obj in page.get('Contents', []):
            key = obj['Key']
            file_obj = s3.get_object(Bucket=bucket_name, Key=key)
            file_content = file_obj['Body'].read()
            files.append({'key': key, 'content': file_content})

    with ThreadPoolExecutor(len(files)) as executor:
        futures = [executor.submit(load_single_document, file['key'], file['content']) for file in files]
        documents = [future.result() for future in as_completed(futures)]
    
    return [doc for doc in documents if doc is not None]

def split_documents(documents: list[LangDocument]) -> tuple[list[LangDocument], list[LangDocument]]:
    text_docs, python_docs = [], []
    for doc in documents:
        if doc is not None:
            try:
                file_extension = os.path.splitext(doc.metadata["source"])[1]
            except:
                document = doc[0]
                metadata = document.metadata
                file_extension = os.path.splitext(metadata.get('source', ''))[1]
            if file_extension == ".py":
                python_docs.append(doc)
            else:
                text_docs.append(doc)
    return text_docs, python_docs

def main(device_type, documents):
    text_documents, python_documents = split_documents(documents)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    python_splitter = RecursiveCharacterTextSplitter.from_language(
        language=Language.PYTHON, chunk_size=880, chunk_overlap=200
    )
    
    texts = text_splitter.split_documents(text_documents)
    texts.extend(python_splitter.split_documents(python_documents))
    logging.info(f"Loaded {len(documents)} documents from S3")
    logging.info(f"Split into {len(texts)} chunks of text")
    embeddings = get_embeddings(device_type)
    
    start_time = time.time()
    vector_store = PineconeVectorStore.from_documents(
        texts,
        index_name=re.sub(r'[^a-z0-9-]+', '-', documents[0].metadata["source"].split('/')[0].lower()),
        embedding=embeddings
    )
    end_time = time.time()
    time_taken = end_time - start_time
    print(f"Time taken to process and store embeddings: {time_taken:.2f} seconds")
    print('after Pinecone')

if __name__ == "__main__":
    # For testing this script seperately 
    bucket_name = os.environ.get('S3_BUCKET_NAME')
    folder = SOURCE_DIRECTORY
    documents = load_documents_from_s3(bucket_name, folder)
    main('cpu', documents)
